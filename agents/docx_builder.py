"""
Gap Step 31: DOCX Rebuild Agent
Rebuilds a standardized .docx file from normalized JSON content.
Python controls ALL Word formatting via python-docx.
OpenAI never controls styling — only content normalization (Step 30).
Output files are written to the docx_output/ folder.
"""

import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.docx_normalizer import NormalizedDocument

# Output folder relative to project root
OUTPUT_DIR = Path(__file__).resolve().parent.parent / "docx_output"


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def rebuild_document(normalized: NormalizedDocument, output_path: str = None) -> str:
    """
    Rebuild a clean .docx file from a NormalizedDocument.

    Args:
        normalized: NormalizedDocument returned by docx_normalizer.normalize_document()
        output_path: Optional full output path. Defaults to docx_output/<file_name>

    Returns:
        Absolute path to the saved .docx file.
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if output_path is None:
        stem = Path(normalized.file_name).stem
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        output_path = str(OUTPUT_DIR / f"{stem}_standardized_{timestamp}.docx")

    doc = Document()
    _apply_document_settings(doc)
    _build_content(doc, normalized)
    _add_page_numbers(doc)

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Document-level settings
# ---------------------------------------------------------------------------

def _apply_document_settings(doc: Document) -> None:
    """Set margins and default font for the entire document."""
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Default paragraph font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


# ---------------------------------------------------------------------------
# Content builder
# ---------------------------------------------------------------------------

def _build_content(doc: Document, normalized: NormalizedDocument) -> None:
    """Inject all normalized sections into the document with correct styles."""

    for section in normalized.sections:
        name = section.section_name
        content = section.content.strip()

        if name == "Title":
            _add_title(doc, content or normalized.file_name)

        elif name == "Document Control":
            _add_heading1(doc, name)
            if content:
                _add_body(doc, content)
            else:
                _add_placeholder(doc, name)

        elif name == "Revision History":
            _add_heading1(doc, name)
            if content:
                _add_revision_table(doc, content)
            else:
                _add_placeholder(doc, name)

        else:
            _add_heading1(doc, name)
            if content:
                _add_body(doc, content)
            else:
                _add_placeholder(doc, name)

    # Confidence footer note
    _add_confidence_note(doc, normalized)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Word dark blue
    para.space_after = Pt(12)


def _add_heading1(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="Heading 1")
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    para.space_before = Pt(12)
    para.space_after = Pt(4)


def _add_body(doc: Document, text: str) -> None:
    """Add body text, splitting on newlines into separate paragraphs."""
    for line in text.splitlines():
        line = line.strip()
        if line:
            para = doc.add_paragraph(style="Normal")
            para.add_run(line)
            para.space_after = Pt(6)


def _add_placeholder(doc: Document, section_name: str) -> None:
    para = doc.add_paragraph(style="Normal")
    run = para.add_run(f"[{section_name} — not present in source document]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    para.space_after = Pt(6)


def _add_revision_table(doc: Document, content: str) -> None:
    """Render revision history as a simple 3-column table."""
    lines = [l.strip() for l in content.splitlines() if l.strip()]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for cell, label in zip(hdr_cells, ["Version", "Description", "Date"]):
        cell.text = label
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Content rows — each line becomes one row
    for line in lines:
        row_cells = table.add_row().cells
        parts = line.split(" – ", 2)  # e.g. "Version 1.0 – Initial release – Jan 2024"
        for i, part in enumerate(parts[:3]):
            row_cells[i].text = part.strip()

    doc.add_paragraph()  # spacing after table


def _add_confidence_note(doc: Document, normalized: NormalizedDocument) -> None:
    doc.add_paragraph()
    para = doc.add_paragraph(style="Normal")
    note = (
        f"Standardized by AP Automation Agent | "
        f"Confidence: {normalized.overall_confidence:.0%} | "
        f"Review required: {'Yes' if normalized.requires_human_review else 'No'}"
    )
    run = para.add_run(note)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Page numbering
# ---------------------------------------------------------------------------

def _add_page_numbers(doc: Document) -> None:
    """Add 'Page X of Y' to the footer of every section."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run("Page ")
        _insert_field(run, "PAGE")
        para.add_run(" of ")
        run2 = para.add_run()
        _insert_field(run2, "NUMPAGES")


def _insert_field(run, field_type: str) -> None:
    """Insert a Word field (PAGE / NUMPAGES) into a run."""
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.text = f" {field_type} "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


# ---------------------------------------------------------------------------
# CLI test entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys
    from dotenv import load_dotenv
    load_dotenv()

    from agents.docx_extractor import extract_docx
    from agents.docx_normalizer import normalize_document

    if len(sys.argv) < 2:
        print("Usage: python docx_builder.py <path_to_docx>")
        sys.exit(1)

    extracted = extract_docx(sys.argv[1])
    normalized = normalize_document(extracted)
    output = rebuild_document(normalized)
    print(f"Saved: {output}")
