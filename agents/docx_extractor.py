"""
Gap Step 29: Document Type Handler
Extracts structured content from .docx files using python-docx.
Outputs a normalized dict ready for OpenAI normalization (Step 30).
"""

import os
from datetime import datetime
from typing import Optional
from docx import Document
from docx.oxml.ns import qn


def extract_docx(file_path: str) -> dict:
    """
    Extract heading hierarchy, paragraphs, tables, and metadata from a .docx file.

    Args:
        file_path: Absolute path to the .docx file.

    Returns:
        Structured dict with keys: file_name, metadata, headings, sections, tables, raw_paragraphs.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is not a valid .docx.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    if not file_path.lower().endswith(".docx"):
        raise ValueError(f"Expected a .docx file, got: {file_path}")

    doc = Document(file_path)

    metadata = _extract_metadata(doc, file_path)
    headings, sections, raw_paragraphs = _extract_paragraphs(doc)
    tables = _extract_tables(doc)

    return {
        "file_name": os.path.basename(file_path),
        "file_path": file_path,
        "metadata": metadata,
        "headings": headings,
        "sections": sections,
        "tables": tables,
        "raw_paragraphs": raw_paragraphs,
        "extracted_at": datetime.utcnow().isoformat() + "Z",
    }


def _extract_metadata(doc: Document, file_path: str) -> dict:
    """Extract core properties from the document."""
    props = doc.core_properties
    return {
        "title": props.title or "",
        "author": props.author or "",
        "created": props.created.isoformat() if props.created else "",
        "modified": props.modified.isoformat() if props.modified else "",
        "subject": props.subject or "",
        "description": getattr(props, "description", "") or "",
        "file_size_bytes": os.path.getsize(file_path),
    }


def _extract_paragraphs(doc: Document) -> tuple[list, list, list]:
    """
    Walk all paragraphs, building:
    - headings: flat list of {level, text} dicts
    - sections: heading-grouped content blocks
    - raw_paragraphs: every paragraph as {style, text}
    """
    headings = []
    raw_paragraphs = []
    sections = []

    current_section = None

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()

        raw_paragraphs.append({
            "style": style_name,
            "text": text,
        })

        if style_name.startswith("Heading"):
            # Parse heading level, default to 1 if not parseable
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1

            heading_entry = {"level": level, "text": text}
            headings.append(heading_entry)

            # Save previous section if it exists
            if current_section is not None:
                sections.append(current_section)

            # Start a new section
            current_section = {
                "heading": text,
                "heading_level": level,
                "content": [],
            }

        elif style_name == "Title":
            # Treat Title style as a level-0 heading
            headings.append({"level": 0, "text": text})
            if current_section is not None:
                sections.append(current_section)
            current_section = {
                "heading": text,
                "heading_level": 0,
                "content": [],
            }

        else:
            # Body paragraph: attach to current section or a default section
            if text:
                if current_section is None:
                    current_section = {
                        "heading": "",
                        "heading_level": None,
                        "content": [],
                    }
                current_section["content"].append(text)

    # Append the last open section
    if current_section is not None:
        sections.append(current_section)

    return headings, sections, raw_paragraphs


def _extract_tables(doc: Document) -> list:
    """
    Extract all tables as lists of row dicts.
    Each table is a list of rows; each row is a list of cell strings.
    """
    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        tables.append({
            "table_index": table_index,
            "row_count": len(rows),
            "column_count": len(rows[0]) if rows else 0,
            "rows": rows,
        })
    return tables


if __name__ == "__main__":
    import sys
    import json

    if len(sys.argv) < 2:
        print("Usage: python docx_extractor.py <path_to_docx>")
        sys.exit(1)

    result = extract_docx(sys.argv[1])
    print(json.dumps(result, indent=2, default=str))
